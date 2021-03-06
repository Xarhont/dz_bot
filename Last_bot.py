import telebot
from telebot import apihelper, types
from bd_config import *
from datetime import datetime
from telebot.types import InlineKeyboardButton, InlineKeyboardMarkup
from docx import Document
import docx
import shutil
from googtab import svobodn_string, zapis_rez
from my_func import *
import time

# apihelper.proxy = {'https': 'socks5://190737618:TsT9nZls@orbtl.s5.opennetwork.cc:999'}  # работал до 29 мая
# apihelper.proxy = {'https': 'socks5://142.93.170.92:1080'}  # 30 июня - 1
# apihelper.proxy = {'https': 'socks5://165.22.65.160:19488'}  # 30 июня - 2
# apihelper.proxy = {'https': 'socks5://78.46.218.20:12041'}  # 30 июня - 3
# apihelper.proxy = {'https': 'socks5://1rje3TFpVJ:Er2GduoOmw@45.92.172.55:55276'}  # куплена до 8 сентября
my_bot = telebot.TeleBot('1245059539:AAEL9lRvA46urwPZmWessOONPgP920cjaTg', num_threads=4)
init_db()

##для учителя
# переменные
teacher = 646951760
db_append_status = ''  # статут добавления нового задания
new_example: Example  # объект новое задание
db_edit_status = ''  # статут изменения задания
edit_example: Example  # объект изменяемого задания
new_multi_dz_theme: Multi_dz_theme  # объект тема для мульти дз
new_multi_dz_themes = {}  # справочник тем и кол-ва заданий для мульти дз С ПУСТЫМИ ТЕМАМИ
new_multi_dz = ''  # справочник тем и кол-ва заданий для мульти дз БЕЗ ПУСТЫХ ТЕМ
new_multi_dz_status = ''
# главная клавиатура
main_page_markup = types.ReplyKeyboardMarkup(True)
main_page_markup.row('База заданий', 'Работа с ДЗ')
# клавиатура работы с БД
main_bd_markup = types.ReplyKeyboardMarkup(True)
main_bd_markup.row('Темы', 'Дополнить БД ОГЭ', 'Изменить задание')
main_bd_markup.row('На главную')
# клавиатура работы с ДЗ
main_dz_markup = types.ReplyKeyboardMarkup(True)
main_dz_markup.row('Задать ДЗ', 'Проверить ДЗ', 'Удалить ДЗ')
main_dz_markup.row('На главную')

##для студента
# главная клавиатура
main_student_markup = types.ReplyKeyboardMarkup(True)
main_student_markup.row('На главную')


@my_bot.message_handler(commands=['start'])
def start_message(message):
    user_id = str(message.from_user.id)
    # интерфейс учителя---------------------------------------------------
    if user_id == teacher:
        my_bot.send_message(user_id, 'Привет, создатель!', reply_markup=main_page_markup)
        my_bot.send_message(user_id, 'Чем займемся?')
    # интерфейс пользователя---------------------------------------------
    else:
        if UserTab.select().where(UserTab.teleg_id == user_id).count() == 0:
            UserTab.create(teleg_id=user_id,
                           name='',
                           klass='',
                           cur_selftest_1t='',
                           reg_status='Нет фам',
                           reg_date=datetime.now(),
                           status='',
                           cur_multitest='нет')
        if UserTab.get(teleg_id=user_id).reg_status == 'Нет фам':
            my_bot.send_message(user_id, 'Привет🖐, давай тебя зарегистрируем в системе.')
            my_bot.send_message(user_id, 'Введи свою Фамилию')
        if UserTab.get(teleg_id=user_id).reg_status == 'Нет имени':
            my_bot.send_message(user_id, 'Введи своё Имя')
        if UserTab.get(teleg_id=user_id).reg_status == 'Выполнена':
            my_bot.send_message(user_id, f'Привет🖐, {UserTab.get(teleg_id=user_id).name}',
                                reply_markup=main_student_markup)
            user_function(message)


@my_bot.message_handler(content_types=['text'])
def main(message):
    user_id = message.chat.id

    # для учителя--------------------------------------------------------------------
    global db_append_status
    global new_example
    global new_dz_status
    if user_id == teacher:
        if message.text == 'На главную':
            clearstatus()
            my_bot.send_message(user_id, 'Чем займемся?', reply_markup=main_page_markup)

        if message.text == 'База заданий':
            my_bot.send_message(user_id, 'Выбери действие 👇', reply_markup=main_bd_markup)
        if message.text == 'Дополнить БД ОГЭ':  # запуск функции добавления заданий (шаг1)
            clearstatus()
            OGE_DB_APPEND1(user_id)
        if db_append_status == 'ожидание ответа':  # получение отвена на новое задание
            new_example.answer = message.text.upper()
            new_example_added_to_bd(new_example)
            my_bot.send_message(user_id, 'Новое задание добавлено')
            db_append_status = ''
            # возможность сразу добавить еще одно задание по этой теме
            more_add_ex = types.InlineKeyboardMarkup()
            more_add_ex.row(types.InlineKeyboardButton(text='Да',
                                                       callback_data=f'example add_{Theme.get(id=str(new_example.theme)).id}'))
            my_bot.send_message(user_id, 'Добавить еще одно задание в эту тему?', reply_markup=more_add_ex)

        if db_edit_status == 'ожидание нового ответа':  # запись нового овтета на задание в бд
            edit_example.answer = message.text
            TestExample.update({TestExample.answer: edit_example.answer}).where(
                TestExample.id == edit_example.id).execute()
            my_bot.send_message(user_id, 'Ответ на задание изменен')
            clearstatus()
        elif message.text == 'Изменить задание':  # запуск функции изменения задания
            clearstatus()
            OGE_DB_EDIT1(user_id)
        elif message.text == 'Задать ДЗ':
            clearstatus()
            create_multi_dz(message)


        elif message.text == 'Назад':
            my_bot.send_message(user_id, 'Выбери действие 👇', reply_markup=main_dz_markup)
        elif message.text == 'Удалить ДЗ':
            clearstatus()
            dz_delete1(message)
        elif message.text == 'Работа с ДЗ':
            my_bot.send_message(user_id, 'Выбери действие 👇', reply_markup=main_dz_markup)

        elif message.text == 'Проверить ДЗ':
            clearstatus()
            multidz_check1(message)

        elif new_multi_dz_status == 'жду названия мультидз':
            kl, zd = new_multi_dz.split('/')
            MultiDzTable.create(klass=Klass.get(name=kl),
                                name=message.text, zadanie=zd, date_create=datetime.now())
            my_bot.send_message(user_id, f'МультиДЗ для {kl} создано 👍')
            clearstatus()
        elif message.text == 'Темы':
            themekey(my_bot, user_id)

        elif db_edit_status == 'жду новой темы':
            Theme.create(name=message.text, archive='NO')
            my_bot.send_message(user_id, 'Тема добавлена')
            clearstatus()

        elif db_edit_status.split('_')[0] == 'Жду нового названия темы для':
            Theme.update(name=message.text).where(Theme.id == int(db_edit_status.split('_')[1])).execute()
            my_bot.send_message(user_id, 'Название темы изменено')
    # для студента--------------------------------------------------------------------
    else:
        if UserTab.select().where(UserTab.teleg_id == user_id).count() == 0:
            my_bot.send_message(user_id, 'Вы не зарегистрированы. Напишите команду /start')
        else:
            if UserTab.get(teleg_id=user_id).reg_status == 'Нет фам':  # если нет фам, записываем и даем выбрать класс
                reg_fam(message)
            elif UserTab.get(
                    teleg_id=user_id).reg_status == 'Нет имени':  # если нет фам, записываем и даем выбрать класс
                reg_name(message)

            elif UserTab.get(
                    teleg_id=user_id).reg_status == 'Выполнена' and message.text == 'На главную':  # если юзер зареган, предлагаем ему выбор функций
                UserTab.update({UserTab.status: ''}).where(UserTab.teleg_id == user_id).execute()
                user_function(message)
                print(f'{datetime.now()} 👉🏻 {UserTab.get(teleg_id=user_id).name} 👉🏻 {message.text}')
            elif UserTab.get(teleg_id=user_id).status == 'мультитест старт':
                start_multitest(message)
            elif UserTab.get(teleg_id=user_id).status == 'мультитест в процессе':
                print(f'{datetime.now()} 👉🏻 {UserTab.get(teleg_id=user_id).name} 👉🏻 {message.text}')
                do_multitest(message)
            my_bot.send_message(user_id, f'☝☝☝',
                                reply_markup=main_student_markup)


##для учителя------------------------------------------------------------------------

# удалени ДЗ, выбор дз из общего списка для удаления
def dz_delete1(message):
    user_id = message.chat.id
    dz_del_key = InlineKeyboardMarkup()
    for dz in MultiDzTable.select()[::-1]:
        dz_del_key.row(types.InlineKeyboardButton(
            text=f'{dz.name} для {dz.klass.name} от {dz.date_create.strftime("%H:%M - %d.%m")}',
            callback_data=f'del dz_{dz.id}'))
    my_bot.send_message(user_id, 'Какое ДЗ удалить?', reply_markup=dz_del_key)


@my_bot.callback_query_handler(func=lambda call: call.data == 'theme_add')
def theme_add(call):
    user_id = call.message.chat.id
    global db_edit_status
    db_edit_status = 'жду новой темы'
    my_bot.send_message(user_id, 'Напишите название новой темы 👇')


@my_bot.callback_query_handler(func=lambda call: call.data.split('_')[0] == 'del dz')
def dz_delete2(call):
    user_id = call.message.chat.id
    del_dz_keyboard = types.InlineKeyboardMarkup()
    del_dz_keyboard.row(types.InlineKeyboardButton(text='Да',
                                                   callback_data=f"del dz1_{call.data.split('_')[1]}"))
    my_bot.send_message(user_id, 'Уверены, что хотите удалить это ДЗ?', reply_markup=del_dz_keyboard)


@my_bot.callback_query_handler(func=lambda call: call.data.split('_')[0] == 'del dz1')
def dz_delete3(call):
    # MultiDzTable.get(id=(call.data.split('_')[1])).delete_instance()
    # my_bot.answer_callback_query(call.id, text="ДЗ удалено")
    ### ver 2.0
    for multitest in MultiDzTable.get(id=(call.data.split('_')[1])).tests:
        for ex in multitest.tests_ex:
            SelfTest_1t_ex.get(id=ex.id).delete_instance()
        MultiTest.get(id=multitest.id).delete_instance()
    MultiDzTable.get(id=(call.data.split('_')[1])).delete_instance()
    my_bot.answer_callback_query(call.id, text="ДЗ удалено")


# начало создания мульти-ДЗ, создание справочника тем и кол-ва заданий
def create_multi_dz(message):
    user_id = message.chat.id
    global new_multi_dz_themes
    global new_multi_dz_theme
    for theme in Theme.select().where(Theme.archive == 'NO'):
        new_multi_dz_themes[theme.id] = Multi_dz_theme(tema=theme.name, id=theme.id, active='no', count=0)
    create_multi_dz_key(message)


# формирование стартовой клавиатуры выбора заданий
def create_multi_dz_key(message):
    global new_multi_dz_themes
    user_id = message.chat.id
    multi_dz_theme = types.InlineKeyboardMarkup()
    for theme in new_multi_dz_themes.values():
        text = theme.tema
        if theme.active == 'yes':
            text += '✅'
        multi_dz_theme.row(types.InlineKeyboardButton(text=text,
                                                      callback_data=f"append to multidz_{theme.id}"))
    multi_dz_theme.row(types.InlineKeyboardButton(text='Далее ➡️',
                                                  callback_data=f"create_multi_dz2"))
    my_bot.send_message(user_id, 'Выберите темы для мульти ДЗ 👇', reply_markup=multi_dz_theme)


# обновление клавиатуры с отметкой выбранных тем
@my_bot.callback_query_handler(func=lambda call: call.data.split('_')[0] == 'append to multidz')
def create_multi_dz2(call):
    global new_multi_dz_themes
    user_id = call.message.chat.id
    id = int(call.data.split('_')[1])
    if new_multi_dz_themes[id].active == 'no':
        new_multi_dz_themes[id].active = 'yes'
    else:
        new_multi_dz_themes[id].active = 'no'
        new_multi_dz_themes[id].count = 0
    multi_dz_keyboard(call)  # вызов клавиатуры с выбором тем и кол-ва заданий


# обновление клавиатуры с отметкой выбранных тем
@my_bot.callback_query_handler(func=lambda call: call.data.split('_')[0] == 'multdz count')
def create_multi_dz2(call):
    global new_multi_dz_themes
    user_id = call.message.chat.id
    new_multi_dz_themes[int(call.data.split('_')[1])].count = int(call.data.split('_')[2])
    multi_dz_keyboard(call)


def multi_dz_keyboard(call):
    multi_dz_theme1 = types.InlineKeyboardMarkup()
    for theme in new_multi_dz_themes.values():
        text = theme.tema
        if theme.active == 'yes':
            text += '✅'
            if theme.count > 0:
                text += f'👉 {theme.count}'
        multi_dz_theme1.row(types.InlineKeyboardButton(text=text,
                                                       callback_data=f"append to multidz_{theme.id}"))

        # кнопки кол-ва заданий
        if theme.active == 'yes' and theme.count == 0:
            multi_dz_theme1.row(types.InlineKeyboardButton(text=1,
                                                           callback_data=f"multdz count_{theme.id}_1"),
                                types.InlineKeyboardButton(text=2,
                                                           callback_data=f"multdz count_{theme.id}_2"),
                                types.InlineKeyboardButton(text=3,
                                                           callback_data=f"multdz count_{theme.id}_3"),
                                types.InlineKeyboardButton(text=4,
                                                           callback_data=f"multdz count_{theme.id}_4"),
                                types.InlineKeyboardButton(text=5,
                                                           callback_data=f"multdz count_{theme.id}_5"))
            multi_dz_theme1.row(types.InlineKeyboardButton(text=6,
                                                           callback_data=f"multdz count_{theme.id}_6"),
                                types.InlineKeyboardButton(text=7,
                                                           callback_data=f"multdz count_{theme.id}_7"),
                                types.InlineKeyboardButton(text=8,
                                                           callback_data=f"multdz count_{theme.id}_8"),
                                types.InlineKeyboardButton(text=9,
                                                           callback_data=f"multdz count_{theme.id}_9"),
                                types.InlineKeyboardButton(text=10,
                                                           callback_data=f"multdz count_{theme.id}_10"))

    multi_dz_theme1.row(types.InlineKeyboardButton(text='Далее ➡️',
                                                   callback_data=f"create_multi_dz2"))

    my_bot.edit_message_text(chat_id=call.message.chat.id, message_id=call.message.message_id,
                             text='Выберите темы для мульти ДЗ 👇',
                             reply_markup=multi_dz_theme1)


# создание строки с темами и кол-вом заданийб изменение статуса мультидз
@my_bot.callback_query_handler(func=lambda call: call.data == 'create_multi_dz2')
def create_multi_dz2(call):
    global new_multi_dz_themes
    global new_multi_dz
    global new_multi_dz_status
    z = ''
    user_id = call.message.chat.id
    for theme in new_multi_dz_themes.values():
        if theme.active == 'yes' and theme.count > 0:
            if new_multi_dz != '':
                z = ';'
            new_multi_dz += f'{z}{theme.tema}_{theme.count}'

    # формирование клавы выбора класса
    user_id = call.message.chat.id
    dz_klass_key = InlineKeyboardMarkup()
    # все классы из 7 параллели
    p7 = []
    for kl in Parallel.get(name='7').classes:
        p7.append(InlineKeyboardButton(text=kl.name, callback_data=f"create multi dz3_{kl.name}"))
    dz_klass_key.row(*p7)
    # 8 параллель
    p8 = []
    for kl in Parallel.get(name='8').classes:
        p8.append(InlineKeyboardButton(text=kl.name, callback_data=f"create multi dz3_{kl.name}"))
    dz_klass_key.row(*p8)
    # 9 параллель
    p9 = []
    for kl in Parallel.get(name='9').classes:
        p9.append(InlineKeyboardButton(text=kl.name, callback_data=f"create multi dz3_{kl.name}"))
    dz_klass_key.row(*p9)
    # 10 параллель
    p10 = []
    for kl in Parallel.get(name='10').classes:
        p10.append(InlineKeyboardButton(text=kl.name, callback_data=f"create multi dz3_{kl.name}"))
    dz_klass_key.row(*p10)
    # 11 параллель
    p11 = []
    for kl in Parallel.get(name='11').classes:
        p11.append(InlineKeyboardButton(text=kl.name, callback_data=f"create multi dz3_{kl.name}"))
    dz_klass_key.row(*p11)
    # вывод на кран клавы с классами
    my_bot.send_message(user_id, 'Для какого класса задать МультиДЗ? 👇', reply_markup=dz_klass_key)


# создание строки с темами и кол-вом заданийб изменение статуса мультидз
@my_bot.callback_query_handler(func=lambda call: call.data.split('_')[0] == 'create multi dz3')
def create_multi_dz2(call):
    global new_multi_dz_themes
    global new_multi_dz
    global new_multi_dz_status
    user_id = call.message.chat.id
    new_multi_dz = f"{call.data.split('_')[1]}/{new_multi_dz}"
    new_multi_dz_status = 'жду названия мультидз'
    my_bot.send_message(user_id, "Введите название для мультиДЗ 👇")


# начало проверки МультиДЗ, выбор класса
def multidz_check1(message):
    user_id = message.chat.id
    dz_klass_key = InlineKeyboardMarkup()
    # все классы из 7 параллели
    p7 = []
    for kl in Parallel.get(name='7').classes:
        p7.append(InlineKeyboardButton(text=kl.name, callback_data=f"multidz check_{kl.name}"))
    dz_klass_key.row(*p7)
    # 8 параллель
    p8 = []
    for kl in Parallel.get(name='8').classes:
        p8.append(InlineKeyboardButton(text=kl.name, callback_data=f"multidz check_{kl.name}"))
    dz_klass_key.row(*p8)
    # 9 параллель
    p9 = []
    for kl in Parallel.get(name='9').classes:
        p9.append(InlineKeyboardButton(text=kl.name, callback_data=f"multidz check_{kl.name}"))
    dz_klass_key.row(*p9)
    # 10 параллель
    p10 = []
    for kl in Parallel.get(name='10').classes:
        p10.append(InlineKeyboardButton(text=kl.name, callback_data=f"multidz check_{kl.name}"))
    dz_klass_key.row(*p10)
    # 11 параллель
    p11 = []
    for kl in Parallel.get(name='11').classes:
        p11.append(InlineKeyboardButton(text=kl.name, callback_data=f"multidz check_{kl.name}"))
    dz_klass_key.row(*p11)
    # вывод на кран клавы с классами
    my_bot.send_message(user_id, 'Какой класс открыть?', reply_markup=dz_klass_key)


# Продолжение проверки ДЗ, выбор ДЗ
@my_bot.callback_query_handler(func=lambda call: call.data.split('_')[0] == 'multidz check')
def dz_check2(call):
    user_id = call.message.chat.id
    # вывод всех дз по параллели
    dz_klass_check = types.InlineKeyboardMarkup()
    for dz in Klass.get(name=call.data.split('_')[1]).multidz_po_klassu:
        dz_klass_check.row(types.InlineKeyboardButton(
            text=f'от {dz.date_create.strftime("%d.%m")} 👉 {dz.name}',
            callback_data=f"open multidz_{dz.id}"),
            types.InlineKeyboardButton(text='Выгрузить в файл', callback_data=f'download multidz_{dz.id}'))
    my_bot.send_message(user_id, 'Какое ДЗ открыть?', reply_markup=dz_klass_check)


##выгрузка 2.0
@my_bot.callback_query_handler(func=lambda call: call.data.split('_')[0] == 'download multidz')
def dz_download1(call):
    user_id = call.message.chat.id
    dz = MultiDzTable.get(id=call.data.split('_')[1])
    dz_otchet = dz.tests.order_by(MultiTest.user.name)
    shutil.copy2('shablon.docx', 'otchet_dz.docx')
    doc = Document('otchet_dz.docx')
    per = 0
    nn = my_bot.send_message(user_id, f'выгрузка {per}/{dz_otchet.count()}')

    tems = {}
    for tema in dz.zadanie.split(';'):
        tems[tema.split('_')[0]] = {'count': 0, 'true': 0}

    # for user in dz.klass.users.order_by(UserTab.name):
    #     print(user.name)
    doc.add_heading(f'{dz.name} от {dz.date_create.strftime("%H:%M - %d.%m")}', 1)
    table = doc.add_table(rows=len(dz.zadanie.split(';')), cols=2)
    table.style = 'Table Grid'
    per = 0

    for user in MultiDzTable.get(id=call.data.split('_')[1]).klass.users.order_by(UserTab.name):
        for userdz in MultiDzTable.get(id=call.data.split('_')[1]).tests.select().where(MultiTest.user == user):

            if userdz.date_finish == '':
                findate = 'незакончено'
            else:
                findate = userdz.date_finish.strftime("%H:%M - %d.%m")
            doc.add_heading(
                f'{str(userdz.user.name).ljust(20, " ")} верно {str(userdz.right_count).rjust(2, " ")} из {str(userdz.ex_count).rjust(2, " ")} выполнение {userdz.date_start.strftime("%H:%M")}-{findate}',
                1)
            for test in userdz.tests_ex:
                if test.right == 'True':
                    text1 = 'Верно ✅'
                    tems[f'{test.test_ex_id.theme.name}']['count'] += 1
                    tems[f'{test.test_ex_id.theme.name}']['true'] += 1
                else:
                    text1 = 'Неверно ❌'
                    tems[f'{test.test_ex_id.theme.name}']['count'] += 1
                doc.add_heading(text1, 2)
                file_info = my_bot.get_file(test.test_ex_id.photo)
                downloadfile = my_bot.download_file(file_info.file_path)
                src = 'D:/Oge test bot 2.0/documents/' + '123.jpg'
                with open(src, 'wb') as new_file:
                    new_file.write(downloadfile)
                doc.add_picture('D:/Oge test bot 2.0/documents/123.jpg', width=docx.shared.Cm(10))
            per += 1
            my_bot.edit_message_text(chat_id=user_id, message_id=nn.message_id,
                                     text=f'выгрузка {per}/{dz_otchet.count()}')
    row = 0
    for tem in tems:
        table.cell(row, 0).text = str(tem)

        table.cell(row, 1).text = str(round(tems.get(tem).get('true') / tems.get(tem).get('count'), 3) * 100)

        row += 1

    doc.save('otchet_dz.docx')
    f = open('otchet_dz.docx', "rb")
    my_bot.send_document(user_id, f)
    print('выгрузка закончена')


# Продолжение проверки ДЗ, выбор конкретного теста юзера
@my_bot.callback_query_handler(func=lambda call: call.data.split('_')[0] == 'open multidz')
def dz_check3(call):
    user_id = call.message.chat.id
    dz_user_check = types.InlineKeyboardMarkup()

    for user in MultiDzTable.get(id=call.data.split('_')[1]).klass.users.order_by(UserTab.name):
        for dz in MultiDzTable.get(id=call.data.split('_')[1]).tests.select().where(MultiTest.user == user):
            if dz.date_finish == '':
                findate = 'незакончено'
            else:
                findate = dz.date_finish.strftime("%H:%M - %d.%m")
            dz_user_check.row(types.InlineKeyboardButton(
                text=f'{str(dz.user.name).ljust(25, "_")} верно {str(dz.right_count).rjust(2, " ")} из {str(dz.ex_count).rjust(2, " ")} \n {dz.date_start.strftime("%H:%M - %d.%m")}/{findate}',
                callback_data=f"open user multidz_{dz.id}"))
    my_bot.send_message(user_id, 'Какой тест открыть??', reply_markup=dz_user_check)


@my_bot.callback_query_handler(func=lambda call: call.data.split('_')[0] == 'open user multidz')
def dz_check3(call):
    user_id = call.message.chat.id
    my_bot.send_message(user_id, f'Тест юзера 👨‍🎓 {MultiTest.get(id=call.data.split("_")[1]).user.name}')
    for ex in MultiTest.get(id=call.data.split('_')[1]).tests_ex:
        if ex.right == 'True':
            text = 'Верно ✅'
        else:
            text = 'Неверно ❌'
        my_bot.send_photo(user_id, ex.test_ex_id.photo)
        my_bot.send_message(user_id, f"Ответ юзера: {ex.user_answer} {text}")


@my_bot.callback_query_handler(func=lambda call: call.data.split('_')[0] == 'del theme')
def del_theme(call):
    theme_id = call.data.split('_')[1]
    Theme.update({Theme.archive: 'YES'}).where(Theme.id == theme_id).execute()
    my_bot.answer_callback_query(call.id, text="Тема помещена в архив")


@my_bot.callback_query_handler(func=lambda call: call.data.split('_')[0] == 'red theme')
def red_theme(call):
    global db_edit_status
    user_id = call.message.chat.id
    db_edit_status = f'Жду нового названия темы для_{call.data.split("_")[1]}'
    my_bot.send_message(user_id, 'Введите новое название темы 👇')


@my_bot.callback_query_handler(func=lambda call: call.data.split('_')[0] == 'example add')
def OGE_DB_APPEND3(call):
    global db_append_status
    global new_example
    user_id = call.message.chat.id
    db_append_status = 'ожидание фото'
    new_example = Example(theme=Theme.get(id=call.data.split('_')[1]), photo='', answer='', archive='NO')
    my_bot.send_message(user_id, 'Пришлите фото нового задания')


# изменение существующих заданий, выбор задания (шаг3)
@my_bot.callback_query_handler(func=lambda call: call.data.split('_')[0] == 'theme edit')
def OGE_DB_EDIT2(call):
    global db_edit_status
    global edit_example
    user_id = call.message.chat.id
    db_edit_status = 'выбор задания'
    edit_example = Example(theme=Theme.get(id=call.data.split('_')[1]), photo='', answer='', archive='NO')
    # allTestsOfTheme = TestExample.select().where(
    #     TestExample.theme == Theme.get(id=edit_example.theme) and TestExample.archive == 'NO')
    allTestsOfTheme = TestExample.select().where(
        (TestExample.theme == Theme.get(id=edit_example.theme))).where(TestExample.archive == 'NO')
    for ex in allTestsOfTheme:
        my_bot.send_message(user_id, f'Задание № {ex.id}   Ответ: {ex.answer}')
        my_bot.send_photo(user_id, ex.photo)
        edit_test_keyboard = types.InlineKeyboardMarkup()
        edit_test_keyboard.row(types.InlineKeyboardButton(text='Редактировать',
                                                          callback_data=f'edit test_{ex.id}'),
                               types.InlineKeyboardButton(text='Удалить',
                                                          callback_data=f'delete test_{ex.id}')
                               )
        my_bot.send_message(user_id, 'Выберите действие:', reply_markup=edit_test_keyboard)


# удаление существующего задания
@my_bot.callback_query_handler(func=lambda call: call.data.split('_')[0] == 'delete test')
def OGE_DB_DEL1(call):
    user_id = call.message.chat.id
    del_test_keyboard = types.InlineKeyboardMarkup()
    del_test_keyboard.row(types.InlineKeyboardButton(text='Да',
                                                     callback_data=f"delete test1_{call.data.split('_')[1]}"))
    my_bot.send_message(user_id, 'Уверены, что хотите удалить это задание?', reply_markup=del_test_keyboard)


@my_bot.callback_query_handler(func=lambda call: call.data.split('_')[0] == 'delete test1')
def OGE_DB_DEL2(call):
    # user_id = call.message.chat.id
    # TestExample.get(id=(call.data.split('_')[1])).delete_instance()
    ex_id = call.data.split('_')[1]
    TestExample.update({TestExample.archive: 'YES'}).where(
        TestExample.id == TestExample.get(id=ex_id)).execute()
    my_bot.answer_callback_query(call.id, text="Задание помещено в архив")


# изменение существующего задания выбор нового фото или овтета (шаг4)
@my_bot.callback_query_handler(func=lambda call: call.data.split('_')[0] == 'edit test')
def OGE_DB_EDIT3(call):
    global edit_example
    user_id = call.message.chat.id
    edit_example.id = int(call.data.split('_')[1])
    my_bot.send_message(user_id,
                        f'Редактируем задание № {edit_example.id} 👈🏻 Ответ: {TestExample.get(id=edit_example.id).answer}')
    my_bot.send_photo(user_id, TestExample.get(id=edit_example.id).photo)
    edit_example_keyboard = types.InlineKeyboardMarkup()
    edit_example_keyboard.row(types.InlineKeyboardButton(text=' 📷 Фото ', callback_data='edit photo'))
    edit_example_keyboard.row(types.InlineKeyboardButton(text='📕 Ответ', callback_data='edit answer'))
    my_bot.send_message(user_id, 'Что изменяем в задании?', reply_markup=edit_example_keyboard)


# изменение фото в существующем задании
@my_bot.callback_query_handler(func=lambda call: call.data.split('_')[0] == 'edit photo')
def OGE_DB_EDIT41(call):
    global edit_example
    global db_edit_status
    user_id = call.message.chat.id
    db_edit_status = 'ожидание нового фото'
    my_bot.send_message(user_id, 'Пришлите новое фото задания')


# изменение ответа в существующем задании
@my_bot.callback_query_handler(func=lambda call: call.data.split('_')[0] == 'edit answer')
def OGE_DB_EDIT42(call):
    global edit_example
    global db_edit_status
    user_id = call.message.chat.id
    db_edit_status = 'ожидание нового ответа'
    my_bot.send_message(user_id, 'Пришлите новое ответ на задание')


@my_bot.message_handler(content_types=['photo'])
def main(message):
    user_id = message.chat.id
    if user_id == teacher:
        global db_append_status
        global db_edit_status
        global new_example
        global edit_example
        user_id = message.chat.id
        if db_append_status == 'ожидание фото':  # добавление заданий, шаг4, получили фото, ждем ответа
            new_example.photo = message.photo[len(message.photo) - 1].file_id
            my_bot.send_message(user_id, 'Введите ответ на задание')
            db_append_status = 'ожидание ответа'
        elif db_edit_status == 'ожидание нового фото':  # запись нового фото в бд
            edit_example.photo = message.photo[0].file_id
            TestExample.update({TestExample.photo: edit_example.photo}).where(
                TestExample.id == edit_example.id).execute()
            my_bot.send_message(user_id, 'Фото задания изменено')
            clearstatus()


# добавление заданий, выбор темы для добавления (шаг2)
def OGE_DB_APPEND1(user_id):
    test_themes = Theme.select().where(Theme.archive == 'NO')
    choice_theme_keyboard = types.InlineKeyboardMarkup()
    for theme in test_themes:
        choice_theme_keyboard.row(types.InlineKeyboardButton(text=theme.name,
                                                             callback_data=f'example add_{theme.id}'))
    my_bot.send_message(user_id, 'Какую тему дополним?', reply_markup=choice_theme_keyboard)


# изменение существующих заданий, выбор темы (шаг2)
def OGE_DB_EDIT1(user_id):
    test_themes = Theme.select().where(Theme.archive == 'NO')
    choice_theme_keyboard = types.InlineKeyboardMarkup()
    for theme in test_themes:
        count = theme.examples.count()
        choice_theme_keyboard.row(types.InlineKeyboardButton(text=f'{theme.name} - {count}',
                                                             callback_data=f'theme edit_{theme.id}'))
    my_bot.send_message(user_id, 'Какую тему открыть для изменения?', reply_markup=choice_theme_keyboard)


def clearstatus():
    global db_edit_status
    global db_append_status
    global new_multi_dz_themes
    global new_multi_dz_status
    global new_multi_dz
    new_multi_dz_themes = {}
    db_edit_status = ''
    new_multi_dz_status = ''
    new_multi_dz = ''


##для ученика-------------------------------------------------------------------------
# продолжение регистрации, после ввода фио
def reg_fam(message):
    user_id = message.chat.id
    UserTab.update({UserTab.name: message.text}).where(UserTab.teleg_id == user_id).execute()
    UserTab.update({UserTab.reg_status: 'Нет имени'}).where(UserTab.teleg_id == user_id).execute()
    my_bot.send_message(user_id, 'Напиши свое имя')


def reg_name(message):
    user_id = message.chat.id
    fam = UserTab.get(teleg_id=user_id).name
    UserTab.update({UserTab.name: (fam + ' ' + message.text)}).where(UserTab.teleg_id == user_id).execute()
    UserTab.update({UserTab.reg_status: 'Нет класса'}).where(UserTab.teleg_id == user_id).execute()
    choice_klass = types.InlineKeyboardMarkup()
    for klass in Klass.select():
        choice_klass.row(types.InlineKeyboardButton(text=klass.name,
                                                    callback_data=f'choice klass_{klass.name}'))
    my_bot.send_message(user_id, 'Выбери свой класс', reply_markup=choice_klass)


# завершение регистрации после выбора класса
@my_bot.callback_query_handler(func=lambda call: call.data.split('_')[0] == 'choice klass')
def reg_klass(call):
    user_id = call.message.chat.id
    klass = Klass.get(name=call.data.split('_')[1])
    UserTab.update({UserTab.klass: klass}).where(UserTab.teleg_id == user_id).execute()
    UserTab.update({UserTab.reg_status: 'Выполнена'}).where(UserTab.teleg_id == user_id).execute()
    my_bot.send_message(user_id, f'Регистрация успешно завершена 👍🏻', reply_markup=main_student_markup)


# выбор стартовых действий
def user_function(message):
    user_id = message.chat.id
    choice_func = types.InlineKeyboardMarkup()
    choice_func.row(types.InlineKeyboardButton(text='Выполнить МультиДЗ',
                                               callback_data='check_multi_dz'))
    my_bot.send_message(user_id, 'Чем займемся?', reply_markup=choice_func)


@my_bot.callback_query_handler(func=lambda call: call.data == 'check_multi_dz')
def user_check_multi_dz(call):
    user_id = call.message.chat.id

    if UserTab.select().where(UserTab.teleg_id == user_id).count() > 0:
        print(f'{datetime.now()} 👉🏻 {UserTab.get(teleg_id=user_id).name} 👉🏻 смотрит список ДЗ')
        mass_mult_dz = UserTab.get(teleg_id=user_id).klass.multidz_po_klassu.select()
        my_bot.send_message(user_id, 'Какое Мульти-ДЗ открыть? ❓')
        if len(mass_mult_dz) > 4:
            r = 5
        else:
            r = len(mass_mult_dz) + 1
        for i in range(1, r):
            z1 = mass_mult_dz[-i]
            if MultiDzTable.get(id=z1.id).tests.select().where(
                    MultiTest.user == UserTab.get(teleg_id=user_id)).count() == 0:
                status_dz = 'не выполнено ❌'
                choice_dz = types.InlineKeyboardMarkup()
                choice_dz.row(types.InlineKeyboardButton(
                    text='Выполнить',
                    callback_data=f'choice multidz_{z1.id}')
                )
                my_bot.send_message(user_id, f'{z1.name} - {status_dz}', reply_markup=choice_dz)
            else:
                status_dz = 'не закончено ✍🏻'
                # проверка выполнения дз
                for test in UserTab.get(teleg_id=user_id).multitests.select().where(MultiTest.multidz_id == z1.id):
                    if test.ex_data == '[]':
                        status_dz = 'выполнено ✅'
                if status_dz == 'выполнено ✅':
                    if MultiDzTable.get(id=z1.id).tests.select().where(
                            MultiTest.user == UserTab.get(
                                teleg_id=user_id)).count() < 3:  # max кол-во попыток выполнения ДЗ
                        if MultiDzTable.get(id=z1.id).tests.select().where(
                                MultiTest.user == UserTab.get(teleg_id=user_id))[-1].ex_data == '[]':
                            choice_dz = types.InlineKeyboardMarkup()
                            choice_dz.row(types.InlineKeyboardButton(
                                text='Переделать 🔄',
                                callback_data=f'choice multidz_{z1.id}'))
                            my_bot.send_message(user_id, f'➡️{z1.name} - {status_dz}', reply_markup=choice_dz)
                        else:
                            tid = MultiDzTable.get(id=z1.id).tests.select().where(
                                MultiTest.user == UserTab.get(teleg_id=user_id))[-1].id
                            choice_dz = types.InlineKeyboardMarkup()
                            choice_dz.row(types.InlineKeyboardButton(
                                text='Переделать 🔄',
                                callback_data=f'choice multidz_{z1.id}'),
                                types.InlineKeyboardButton(
                                    text='Продолжить последнюю попытку ',
                                    callback_data=f'continue multidz_{tid}')
                            )
                            my_bot.send_message(user_id, f'➡️{z1.name} - {status_dz}', reply_markup=choice_dz)
                    else:
                        my_bot.send_message(user_id, f'➡️{z1.name} - {status_dz}')
                else:
                    status_dz = 'не закончено ✍🏻'
                    tid = MultiDzTable.get(id=z1.id).tests.select().where(
                        MultiTest.user == UserTab.get(teleg_id=user_id))[-1].id
                    choice_dz = types.InlineKeyboardMarkup()
                    choice_dz.row(types.InlineKeyboardButton(
                        text='Продолжить последнюю попытку',
                        callback_data=f'continue multidz_{tid}'))
                    my_bot.send_message(user_id, f'➡️{z1.name} - {status_dz}', reply_markup=choice_dz)
            # my_bot.send_message(user_id, f'{z1.name} - {status_dz}', reply_markup=choice_dz)
    else:
        my_bot.send_message(user_id, 'Прости, я вчера потерял память. Напиши в чат /start , чтобы зарегистрироваться')
        print(f'Незарегистрированный пользователь {user_id} попробовал открыть список ДЗ')

@my_bot.callback_query_handler(func=lambda call: call.data.split('_')[0] == 'continue multidz')
def continue_multidz1(call):
    user_id = call.message.chat.id
    UserTab.update({UserTab.cur_multitest: int(call.data.split('_')[1])}).where(UserTab.teleg_id == user_id).execute()
    start_multitest(call.message)


# создание мультитеста дз по шаблону и запуск теста
@my_bot.callback_query_handler(func=lambda call: call.data.split('_')[0] == 'choice multidz')
def create_multi_dz1(call):
    dz = MultiDzTable.get(id=call.data.split('_')[1])
    user_id = call.message.chat.id
    if UserTab.select().where(UserTab.teleg_id == user_id).count() > 0:
        tid_dz = MultiTest.create(multidz_id=dz.id,
                                  user=UserTab.get(teleg_id=user_id),
                                  ex_data=gen_numex_multi_dz(dz.zadanie),
                                  ex_count=multidz_count_sum(dz.zadanie),
                                  done_ex_count=0,
                                  right_count=0,
                                  date_start=datetime.now(),
                                  date_finish='').id
        UserTab.update({UserTab.cur_multitest: tid_dz}).where(UserTab.teleg_id == user_id).execute()
        UserTab.update({UserTab.status: 'мультитест старт'}).where(UserTab.teleg_id == user_id).execute()
        print(f'{datetime.now()} 👉🏻 {UserTab.get(teleg_id=user_id).name} 👉🏻 начал выполнение теста')
        start_multitest(call.message)
    else:
        my_bot.send_message(user_id, 'Прости, я на днях потерял память. Нужно снова зарегистрироваться. Напиши в чат /start')
        print(f'Незарегистрированный пользователь {user_id} пытался начать тест')
def start_multitest(message):
    user_id = message.chat.id
    ex_id = MultiTest.get(id=UserTab.get(teleg_id=user_id).cur_multitest).ex_data[1:-1].split(
        ', ')  # срезом убирает квадратные скобки и сплитуем
    my_bot.send_message(user_id, f'Заданий осталось {len(ex_id)}')
    my_bot.send_photo(user_id, TestExample.get(id=ex_id[0]).photo)

    propusk_theme_selftest_1t = types.InlineKeyboardMarkup()
    propusk_theme_selftest_1t.row(types.InlineKeyboardButton(text='Пропустить задание',
                                                             callback_data='propusk_multitest'))

    my_bot.send_message(user_id, 'Введите ваш ответ', reply_markup=propusk_theme_selftest_1t)
    UserTab.update({UserTab.status: 'мультитест в процессе'}).where(UserTab.teleg_id == user_id).execute()


@my_bot.callback_query_handler(func=lambda call: call.data == 'propusk_multitest')
def call_propusk_multitest(call):
    user_id = call.message.chat.id
    if UserTab.select().where(UserTab.teleg_id == user_id).count() > 0:
        if UserTab.get(teleg_id=user_id).cur_multitest != 'нет':
            print(f'Пользователь {UserTab.get(teleg_id=user_id).name} пропустил задание')
            UserTab.update({UserTab.status: 'мультитест старт'}).where(UserTab.teleg_id == user_id).execute()
            multitest_sdvig(user_id)
            start_multitest(call.message)


def do_multitest(message):
    ex_id = []
    user_id = message.chat.id
    ex_id = MultiTest.get(id=UserTab.get(teleg_id=user_id).cur_multitest).ex_data[1:-1].split(', ')

    # пока кол-во выполненых заданий меньше общего кол-ва заданий в тесте
    if MultiTest.get(id=UserTab.get(teleg_id=user_id).cur_multitest).done_ex_count < MultiTest.get(
            id=UserTab.get(teleg_id=user_id).cur_multitest).ex_count:
        # проверяем правильность ответа

        if TestExample.get(id=ex_id[0]).answer == message.text.upper():
            ranswer = 'True'
            print(f'пользователь {UserTab.get(teleg_id=user_id).name} ввел верный ответ')
            MultiTest.update({MultiTest.right_count: MultiTest.right_count + 1}).where(
                MultiTest.id == UserTab.get(teleg_id=user_id).cur_multitest).execute()
        else:
            ranswer = 'False'
            print(f'пользователь {UserTab.get(teleg_id=user_id).name} ввел НЕверный ответ')
        MultiTest.update({MultiTest.done_ex_count: MultiTest.done_ex_count + 1}).where(
            MultiTest.id == UserTab.get(teleg_id=user_id).cur_multitest).execute()
        # создаем запись о задании в БД
        test_id = UserTab.get(teleg_id=user_id).cur_multitest
        SelfTest_1t_ex.create(user=UserTab.get(teleg_id=user_id),
                              test_ex_id=TestExample.get(id=ex_id[0]),
                              user_answer=message.text,
                              right=ranswer,
                              date=datetime.today(),
                              multitest_id=UserTab.get(teleg_id=user_id).cur_multitest)

        # сдвиг номер задания
        ex_id = multitest_sdvig_del(user_id)
        # выдаем следующее задание

        if MultiTest.get(id=UserTab.get(teleg_id=user_id).cur_multitest).done_ex_count < MultiTest.get(
                id=UserTab.get(teleg_id=user_id).cur_multitest).ex_count:
            my_bot.send_message(user_id, f'Заданий осталось {len(ex_id)}')
            my_bot.send_photo(user_id, TestExample.get(id=ex_id[0]).photo)
            propusk_theme_selftest_1t = types.InlineKeyboardMarkup()
            propusk_theme_selftest_1t.row(types.InlineKeyboardButton(text='Пропустить задание',
                                                                     callback_data='propusk_multitest'))

            my_bot.send_message(user_id, 'Введите ваш ответ', reply_markup=propusk_theme_selftest_1t)
        else:
            my_bot.send_message(user_id, 'Тест закончен')
            MultiTest.update(date_finish=datetime.now()).where(
                MultiTest.id == UserTab.get(teleg_id=user_id).cur_multitest).execute()
            my_bot.send_message(user_id,
                                f'Результат {MultiTest.get(id=UserTab.get(teleg_id=user_id).cur_multitest).right_count} из {MultiTest.get(id=UserTab.get(teleg_id=user_id).cur_multitest).ex_count}')
            UserTab.update(cur_multitest='нет').where(UserTab.teleg_id == user_id).execute()
            UserTab.update(status='').where(UserTab.teleg_id == user_id).execute()
            print(f'{datetime.now()} 👉🏻 {UserTab.get(teleg_id=user_id).name} 👉🏻 закончил тест')
    else:
        my_bot.send_message(user_id, 'Тест закончен')
        MultiTest.update(date_finish=datetime.now()).where(
            MultiTest.id == UserTab.get(teleg_id=user_id).cur_multitest).execute()
        my_bot.send_message(user_id,
                            f'Результат {MultiTest.get(id=UserTab.get(teleg_id=user_id).cur_multitest).right_count} из {MultiTest.get(id=UserTab.get(teleg_id=user_id).cur_multitest).ex_count}')
        UserTab.update(cur_multitest='нет').where(UserTab.teleg_id == user_id).execute()
        UserTab.update(status='').where(UserTab.teleg_id == user_id).execute()
        print(f'пользователь {UserTab.get(teleg_id=user_id).name} закончил тест')


# сдвиг номера задания на следующий без удаления
def multitest_sdvig(user_id):
    ex_id = MultiTest.get(id=UserTab.get(teleg_id=user_id).cur_multitest).ex_data[1:-1].split(', ')

    first = ex_id[0]
    for i in range(0, len(ex_id) - 1):
        ex_id[i] = ex_id[i + 1]
    ex_id[len(ex_id) - 1] = first
    for i in range(0, len(ex_id)):
        ex_id[i] = int(ex_id[i])

    MultiTest.update({MultiTest.ex_data: ex_id}).where(
        MultiTest.id == UserTab.get(teleg_id=user_id).cur_multitest).execute()
    return ex_id


# сдвиг номера задания на следующей с удалением
def multitest_sdvig_del(user_id):
    ex_id = MultiTest.get(id=UserTab.get(teleg_id=user_id).cur_multitest).ex_data[1:-1].split(', ')
    ex_id.pop(0)

    for i in range(0, len(ex_id)):
        ex_id[i] = int(ex_id[i])

    MultiTest.update({MultiTest.ex_data: ex_id}).where(
        MultiTest.id == UserTab.get(teleg_id=user_id).cur_multitest).execute()
    return ex_id


er_flag = False
if __name__ == '__main__':
    while True:
        try:
            if er_flag == True:
                my_bot.send_message(teacher, 'У бота ошибка 😢')
            er_flag = False
            my_bot.polling(none_stop=False, interval=0, timeout=20)
        except Exception as e:
            print(e)
            er_flag = True
            time.sleep(5)
